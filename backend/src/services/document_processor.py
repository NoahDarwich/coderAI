"""
Document processing service for parsing PDF, DOCX, and TXT files.

This service extracts text content from uploaded documents using PyMuPDF (for PDFs)
and python-docx (for Word documents).
"""
import io
from typing import BinaryIO

import fitz  # PyMuPDF
from docx import Document as DocxDocument


class DocumentProcessingError(Exception):
    """Raised when document processing fails."""
    pass


def parse_pdf(file: BinaryIO) -> str:
    """
    Parse PDF file and extract text content.

    Uses PyMuPDF (fitz) for fast and accurate text extraction with layout preservation.

    Args:
        file: Binary file object (PDF)

    Returns:
        Extracted text content

    Raises:
        DocumentProcessingError: If PDF parsing fails
    """
    try:
        # Read file content
        file_content = file.read()

        # Open PDF document
        pdf_document = fitz.open(stream=file_content, filetype="pdf")

        # Extract text from all pages
        text_content = []
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            text = page.get_text()

            # Add page separator for multi-page documents
            if page_num > 0:
                text_content.append(f"\n\n--- Page {page_num + 1} ---\n\n")

            text_content.append(text)

        pdf_document.close()

        # Combine all pages
        full_text = "".join(text_content).strip()

        if not full_text:
            raise DocumentProcessingError("PDF contains no extractable text (may be scanned image)")

        return full_text

    except fitz.fitz.FileDataError as e:
        raise DocumentProcessingError(f"Invalid PDF file: {str(e)}")
    except Exception as e:
        raise DocumentProcessingError(f"Failed to parse PDF: {str(e)}")


def parse_docx(file: BinaryIO) -> str:
    """
    Parse DOCX file and extract text content.

    Uses python-docx to extract text from Word documents, preserving paragraph structure.

    Args:
        file: Binary file object (DOCX)

    Returns:
        Extracted text content

    Raises:
        DocumentProcessingError: If DOCX parsing fails
    """
    try:
        # Open DOCX document
        doc = DocxDocument(file)

        # Extract text from all paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Skip empty paragraphs
                paragraphs.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text:
                    paragraphs.append(row_text)

        # Combine all paragraphs
        full_text = "\n\n".join(paragraphs).strip()

        if not full_text:
            raise DocumentProcessingError("DOCX contains no extractable text")

        return full_text

    except Exception as e:
        raise DocumentProcessingError(f"Failed to parse DOCX: {str(e)}")


def parse_txt(file: BinaryIO) -> str:
    """
    Parse plain text file.

    Args:
        file: Binary file object (TXT)

    Returns:
        Text content

    Raises:
        DocumentProcessingError: If TXT parsing fails
    """
    try:
        # Try UTF-8 first
        content = file.read()

        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1 if UTF-8 fails
            try:
                text = content.decode('latin-1')
            except UnicodeDecodeError:
                raise DocumentProcessingError("Unable to decode text file (unsupported encoding)")

        text = text.strip()

        if not text:
            raise DocumentProcessingError("Text file is empty")

        return text

    except DocumentProcessingError:
        raise
    except Exception as e:
        raise DocumentProcessingError(f"Failed to parse text file: {str(e)}")


def parse_html(file: BinaryIO) -> str:
    """
    Parse HTML file and extract text content using BeautifulSoup.

    Args:
        file: Binary file object (HTML)

    Returns:
        Extracted text content

    Raises:
        DocumentProcessingError: If HTML parsing fails
    """
    try:
        from bs4 import BeautifulSoup

        content = file.read()

        # Try UTF-8 first, then latin-1
        try:
            html_text = content.decode("utf-8")
        except UnicodeDecodeError:
            html_text = content.decode("latin-1")

        soup = BeautifulSoup(html_text, "html.parser")

        # Remove script and style elements
        for element in soup(["script", "style", "head"]):
            element.decompose()

        # Get text with newline separation
        text = soup.get_text(separator="\n", strip=True)

        if not text:
            raise DocumentProcessingError("HTML contains no extractable text")

        return text

    except DocumentProcessingError:
        raise
    except Exception as e:
        raise DocumentProcessingError(f"Failed to parse HTML: {str(e)}")


def chunk_text(text: str, max_words: int = 3000, overlap_words: int = 300) -> list[dict]:
    """
    Split text into overlapping chunks for processing large documents.

    Args:
        text: Full document text
        max_words: Maximum words per chunk (default: 3000 words ~= 4000 tokens)
        overlap_words: Number of overlap words between chunks

    Returns:
        List of dicts with keys: text, token_count, overlap_tokens
    """
    words = text.split()

    if len(words) <= max_words:
        return [{"text": text, "token_count": len(words), "overlap_tokens": 0}]

    chunks = []
    start = 0

    while start < len(words):
        end = min(start + max_words, len(words))
        chunk_words = words[start:end]
        overlap = overlap_words if start > 0 else 0
        chunks.append({
            "text": " ".join(chunk_words),
            "token_count": len(chunk_words),
            "overlap_tokens": overlap,
        })
        start = end - overlap_words if end < len(words) else end

    return chunks


async def create_chunks_for_document(db, document, min_word_count: int = 5000) -> int:
    """
    Create document chunks in the database if the document is large enough.

    Args:
        db: AsyncSession
        document: Document model instance
        min_word_count: Minimum word count to trigger chunking

    Returns:
        Number of chunks created (0 if not chunked)
    """
    from src.models.document_chunk import DocumentChunk

    if not document.content or (document.word_count or 0) <= min_word_count:
        return 0

    chunks_data = chunk_text(document.content)
    if len(chunks_data) <= 1:
        return 0

    for i, chunk_data in enumerate(chunks_data):
        chunk = DocumentChunk(
            document_id=document.id,
            chunk_index=i,
            text=chunk_data["text"],
            token_count=chunk_data["token_count"],
            overlap_tokens=chunk_data["overlap_tokens"],
        )
        db.add(chunk)

    document.chunk_count = len(chunks_data)
    return len(chunks_data)


def get_content_preview(text: str, max_length: int = 500) -> str:
    """
    Get preview of document content (first N characters).

    Args:
        text: Full document text
        max_length: Maximum preview length

    Returns:
        Preview text (truncated with "..." if needed)
    """
    if len(text) <= max_length:
        return text

    return text[:max_length] + "..."
